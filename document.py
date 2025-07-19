import streamlit as st
import docx
from docx import Document
import re
import os
from datetime import datetime
from docx.shared import Pt
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter
from docx.enum.text import WD_COLOR_INDEX, WD_LINE_SPACING, WD_ALIGN_PARAGRAPH
from collections import defaultdict
from tkinter import Tk, filedialog
from reportlab.lib import colors
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak, Table, TableStyle
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.enums import TA_CENTER, TA_LEFT

# --- Existing validation functions remain unchanged here ---
# (Keep all the existing functions like check_font_size, check_margins, etc.)

# === NEW CORRECTION FUNCTION ===
def correct_document(doc):
    """Automatically correct document formatting to meet ADF standards"""
    # 1. Correct font sizes and styles
    for para in doc.paragraphs:
        # Set heading and body font sizes
        expected_size = Pt(14) if para.style.name == 'Heading 1' else Pt(12)
        for run in para.runs:
            run.font.size = expected_size
            run.font.name = 'Arial'

    # 2. Adjust margins to 1 inch
    for section in doc.sections:
        section.top_margin = Pt(72)    # 1 inch = 72 points
        section.bottom_margin = Pt(72)
        section.left_margin = Pt(72)
        section.right_margin = Pt(72)

    # 3. Set line spacing to 1.5
    for para in doc.paragraphs:
        para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE

    # 4. Correct caption styles
    for para in doc.paragraphs:
        text = para.text.strip().lower()
        if text.startswith(("table", "figure")):
            for run in para.runs:
                run.font.name = 'Arial'
                run.font.size = Pt(11)

    # 5. Format headers/footers
    for section in doc.sections:
        # Process headers
        for header in [section.header, section.even_page_header, section.first_page_header]:
            if header:
                for para in header.paragraphs:
                    for run in para.runs:
                        run.font.name = 'Arial'
                        run.font.size = Pt(9)
                    # Align page numbers to right
                    if any(char.isdigit() for char in para.text):
                        para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        # Process footers
        for footer in [section.footer, section.even_page_footer, section.first_page_footer]:
            if footer:
                for para in footer.paragraphs:
                    for run in para.runs:
                        run.font.name = 'Arial'
                        run.font.size = Pt(9)

    return doc

# === MODIFIED MAIN FUNCTION ===
def main():
    st.title("ADF Document Checker")
    st.write("Upload your Word document to check for formatting compliance")

    uploaded_file = st.file_uploader("Choose a Word document", type=["docx"])

    if uploaded_file is not None:
        # Save the uploaded file temporarily
        with open("temp.docx", "wb") as f:
            f.write(uploaded_file.getbuffer())
        
        # Load document for validation
        doc = Document("temp.docx")
        
        # --- Existing validation checks remain the same ---
        # (Keep all the validation code and report generation)

        # === NEW DOWNLOAD BUTTON ===
        st.subheader("Document Correction")
        if st.button("Download Corrected Document"):
            # Reload fresh copy for correction
            corrected_doc = Document("temp.docx")
            corrected_doc = correct_document(corrected_doc)
            
            # Save corrected document
            corrected_path = "corrected_ADF_document.docx"
            corrected_doc.save(corrected_path)
            
            # Offer download
            with open(corrected_path, "rb") as f:
                st.download_button(
                    label="Download Corrected Version",
                    data=f,
                    file_name=corrected_path,
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
            
            # Cleanup
            os.remove(corrected_path)

        # Cleanup temporary file
        os.remove("temp.docx")

if __name__ == "__main__":
    main()